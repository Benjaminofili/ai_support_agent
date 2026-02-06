
import os
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

TEST_DOCS_DIR = "test_docs"
DOCX_FILENAME = "verification_test.docx"
PDF_FILENAME = "verification_test.pdf"

def ensure_dir(directory):
    if not os.path.exists(directory):
        os.makedirs(directory)

def create_docx(filepath):
    doc = Document()
    doc.add_heading('Verification Test Document', 0)
    doc.add_paragraph('This is a test document to verify DOCX upload and processing.')
    doc.add_paragraph('The secret verification code is: DOCX-ALPHA-99.')
    doc.add_paragraph('Please answer the question: What is the secret code?')
    doc.save(filepath)
    print(f"Created {filepath}")

def create_pdf(filepath):
    c = canvas.Canvas(filepath, pagesize=letter)
    c.drawString(100, 750, "Verification Test Document (PDF)")
    c.drawString(100, 730, "This is a test document to verify PDF upload and processing.")
    c.drawString(100, 710, "The secret verification code is: PDF-BETA-42.")
    c.drawString(100, 690, "Please answer the question: What is the secret code?")
    c.save()
    print(f"Created {filepath}")

def main():
    ensure_dir(TEST_DOCS_DIR)
    
    docx_path = os.path.join(TEST_DOCS_DIR, DOCX_FILENAME)
    create_docx(docx_path)
    
    pdf_path = os.path.join(TEST_DOCS_DIR, PDF_FILENAME)
    create_pdf(pdf_path)

if __name__ == "__main__":
    main()
