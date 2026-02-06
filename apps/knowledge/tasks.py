"""
Document processing tasks.

This module handles the asynchronous processing of uploaded documents,
including text extraction, chunking, and embedding generation.
"""

import csv
import io
import json
import logging
import os
import tempfile
from abc import ABC, abstractmethod
from typing import List, Optional

from celery import shared_task
from django.conf import settings

from .models import Document, DocumentChunk

logger = logging.getLogger(__name__)

# Constants
CHUNK_SIZE = getattr(settings, "CHUNK_SIZE", 500)
CHUNK_OVERLAP = getattr(settings, "CHUNK_OVERLAP", 50)


# =============================================================================
# TEXT EXTRACTION CLASSES
# =============================================================================


class TextExtractor(ABC):
    """Base class for text extraction from various file formats."""

    file_type: str = "unknown"

    @abstractmethod
    def extract(self, source: str) -> str:
        """
        Extract text from the source.

        Args:
            source: File path or content string

        Returns:
            Extracted text content
        """
        pass

    def log_extraction(self, text: str, details: str = "") -> None:
        """Log extraction results."""
        logger.info(f"{self.file_type} extraction: {len(text)} chars {details}")


class PDFExtractor(TextExtractor):
    """Extract text from PDF files."""

    file_type = "PDF"

    def extract(self, file_path: str) -> str:
        from langchain_community.document_loaders import PyPDFLoader

        loader = PyPDFLoader(file_path)
        pages = loader.load()
        text = "\n\n".join([p.page_content for p in pages])
        self.log_extraction(text, f"({len(pages)} pages)")
        return text


class TextFileExtractor(TextExtractor):
    """Extract text from plain text files."""

    file_type = "TEXT"

    def extract(self, content: bytes) -> str:
        encodings = ["utf-8", "latin-1", "cp1252", "ascii"]

        for encoding in encodings:
            try:
                text = content.decode(encoding)
                self.log_extraction(text, f"(encoding: {encoding})")
                return text
            except UnicodeDecodeError:
                continue

        # Fallback: decode with errors ignored
        text = content.decode("utf-8", errors="ignore")
        self.log_extraction(text, "(with errors ignored)")
        return text


class CSVExtractor(TextExtractor):
    """Extract text from CSV files."""

    file_type = "CSV"

    def extract(self, content: bytes) -> str:
        text_extractor = TextFileExtractor()
        text_content = text_extractor.extract(content)

        try:
            csv_reader = csv.reader(io.StringIO(text_content))
            rows = list(csv_reader)

            if not rows:
                return text_content

            headers = rows[0]
            text_parts = [f"Columns: {', '.join(headers)}", ""]

            for i, row in enumerate(rows[1:], 1):
                row_text = []
                for j, value in enumerate(row):
                    if j < len(headers) and value.strip():
                        row_text.append(f"{headers[j]}: {value}")
                if row_text:
                    text_parts.append(f"Record {i}: {', '.join(row_text)}")

            text = "\n".join(text_parts)
            self.log_extraction(text, f"({len(rows)} rows)")
            return text

        except Exception as e:
            logger.warning(f"CSV parsing error: {e}, using raw text")
            return text_content


class DOCXExtractor(TextExtractor):
    """Extract text from Word documents."""

    file_type = "DOCX"

    def extract(self, file_path: str) -> str:
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(file_path)
            paragraphs = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    ]
                    if row_text:
                        paragraphs.append(" | ".join(row_text))

            text = "\n\n".join(paragraphs)
            self.log_extraction(text, f"({len(paragraphs)} paragraphs)")
            return text

        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise ValueError(f"Could not extract text from DOCX: {e}")


class JSONExtractor(TextExtractor):
    """Extract text from JSON files."""

    file_type = "JSON"

    def extract(self, content: bytes) -> str:
        text_extractor = TextFileExtractor()
        text_content = text_extractor.extract(content)

        try:
            data = json.loads(text_content)
            text_lines = self._flatten_json(data)
            text = "\n".join(text_lines)
            self.log_extraction(text, f"({len(text_lines)} fields)")
            return text

        except json.JSONDecodeError as e:
            logger.warning(f"JSON parsing error: {e}, using raw text")
            return text_content

    def _flatten_json(self, obj, prefix: str = "") -> List[str]:
        """Recursively flatten JSON to readable text."""
        lines = []

        if isinstance(obj, dict):
            for key, value in obj.items():
                new_prefix = f"{prefix}.{key}" if prefix else key
                lines.extend(self._flatten_json(value, new_prefix))
        elif isinstance(obj, list):
            for i, item in enumerate(obj):
                new_prefix = f"{prefix}[{i}]"
                lines.extend(self._flatten_json(item, new_prefix))
        else:
            if obj is not None and str(obj).strip():
                lines.append(f"{prefix}: {obj}")

        return lines


class MarkdownExtractor(TextExtractor):
    """Extract text from Markdown files."""

    file_type = "MARKDOWN"

    def extract(self, content: bytes) -> str:
        text_extractor = TextFileExtractor()
        text = text_extractor.extract(content)
        self.log_extraction(text)
        return text


# =============================================================================
# HELPER FUNCTIONS
# =============================================================================


def get_extractor(source_type: str, extension: str) -> TextExtractor:
    """
    Get the appropriate text extractor based on file type.

    Args:
        source_type: Document source type
        extension: File extension

    Returns:
        TextExtractor: Appropriate extractor instance
    """
    extractors = {
        "pdf": PDFExtractor(),
        ".pdf": PDFExtractor(),
        "docx": DOCXExtractor(),
        ".docx": DOCXExtractor(),
        "csv": CSVExtractor(),
        ".csv": CSVExtractor(),
        "json": JSONExtractor(),
        ".json": JSONExtractor(),
        "markdown": MarkdownExtractor(),
        ".md": MarkdownExtractor(),
        ".markdown": MarkdownExtractor(),
    }

    return (
        extractors.get(source_type.lower())
        or extractors.get(extension.lower())
        or TextFileExtractor()
    )


def get_file_extension(filename: str) -> str:
    """Get lowercase file extension from filename."""
    if not filename:
        return ""
    return os.path.splitext(filename)[1].lower()


def extract_text_from_document(doc: Document) -> str:
    """
    Extract text content from a document.

    Args:
        doc: The Document model instance

    Returns:
        Extracted text content

    Raises:
        ValueError: If no text content could be extracted
    """
    text = ""

    if doc.file:
        filename = doc.file.name if doc.file else ""
        extension = get_file_extension(filename)

        logger.info(f"Processing file: {filename}, extension: {extension}")

        # File-based extraction (PDF, DOCX)
        if extension in [".pdf", ".docx"]:
            suffix = extension
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                doc.file.seek(0)
                for chunk in doc.file.chunks():
                    tmp.write(chunk)
                tmp_path = tmp.name

            try:
                extractor = get_extractor(doc.source_type, extension)
                text = extractor.extract(tmp_path)
            finally:
                os.unlink(tmp_path)

        # Content-based extraction (CSV, JSON, TXT, MD)
        else:
            doc.file.seek(0)
            content = doc.file.read()
            extractor = get_extractor(doc.source_type, extension)
            text = extractor.extract(content)

    # Fallback to raw_content
    if not text or not text.strip():
        if doc.raw_content and doc.raw_content.strip():
            text = doc.raw_content
            logger.info(f"Using raw_content: {len(text)} characters")
        else:
            raise ValueError(
                f"No text content extracted. File: {bool(doc.file)}, "
                f"Raw content: {len(doc.raw_content or '')} chars"
            )

    logger.info(f"Total extracted text: {len(text)} characters")
    return text


def create_chunks(
    text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP
) -> List[str]:
    """
    Split text into overlapping chunks.

    Args:
        text: Text to split
        chunk_size: Maximum characters per chunk
        overlap: Overlap between chunks

    Returns:
        List of text chunks
    """
    from langchain.text_splitter import RecursiveCharacterTextSplitter

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=overlap,
        separators=["\n\n", "\n", ". ", ", ", " ", ""],
    )

    chunks = splitter.split_text(text)
    logger.info(f"Split into {len(chunks)} chunks")

    return chunks


def generate_embedding(text: str) -> List[float]:
    """
    Generate embedding for text.

    Args:
        text: Text to embed

    Returns:
        Embedding vector
    """
    if hasattr(settings, "HUGGINGFACE_API_KEY") and settings.HUGGINGFACE_API_KEY:
        from apps.conversations.huggingface_service import \
            generate_embedding as hf_embed

        return hf_embed(text)

    if hasattr(settings, "OPENAI_API_KEY") and settings.OPENAI_API_KEY:
        from openai import OpenAI

        client = OpenAI(api_key=settings.OPENAI_API_KEY)
        response = client.embeddings.create(model="text-embedding-3-small", input=text)
        return response.data[0].embedding

    # Default to HuggingFace local model
    from apps.conversations.huggingface_service import \
        generate_embedding as hf_embed

    return hf_embed(text)


def save_chunks_with_embeddings(document: Document, chunks: List[str]) -> int:
    """
    Save chunks and generate embeddings for a document.

    Args:
        document: Parent document
        chunks: List of text chunks

    Returns:
        Number of chunks saved
    """
    for i, chunk_text in enumerate(chunks):
        embedding = generate_embedding(chunk_text)

        DocumentChunk.objects.create(
            document=document,
            content=chunk_text,
            chunk_index=i,
            embedding=embedding,
            metadata={
                "source": document.title,
                "chunk": i,
                "total_chunks": len(chunks),
            },
        )

        # Log progress every 5 chunks
        if (i + 1) % 5 == 0 or (i + 1) == len(chunks):
            logger.info(f"Processed {i + 1}/{len(chunks)} chunks")

    return len(chunks)


# =============================================================================
# CELERY TASK
# =============================================================================


@shared_task(bind=True, max_retries=3, time_limit=600, soft_time_limit=540)
def process_document_task(self, document_id: str):
    """
    Process an uploaded document.
    
    Pipeline:
    1. Extract text from the document
    2. Split into chunks
    3. Generate embeddings for each chunk
    4. Save chunks to database
    
    Args:
        document_id: UUID of the document to process
        
    Returns:
        Dict with status and chunk count
        
    Note:
        - Hard timeout: 600 seconds (10 minutes)
        - Soft timeout: 540 seconds (9 minutes)
        - Retries: 3 times with exponential backoff
    """
    try:
        # Get document
        doc = Document.objects.get(id=document_id)
        doc.status = Document.Status.PROCESSING
        doc.error_message = ""
        doc.save()
        
        logger.info(f"Processing document: {doc.title} (ID: {document_id})")
        logger.info(f"Source type: {doc.source_type}, Has file: {bool(doc.file)}")
        
        # Step 1: Extract text
        text = extract_text_from_document(doc)
        
        if not text.strip():
            raise ValueError("Document produced no text content")
        
        # Step 2: Create chunks
        chunks = create_chunks(text)
        
        if not chunks:
            raise ValueError("Document produced no chunks after splitting")
        
        # Step 3 & 4: Generate embeddings and save
        chunk_count = save_chunks_with_embeddings(doc, chunks)
        
        # Update document status
        doc.status = Document.Status.COMPLETED
        doc.chunk_count = chunk_count
        doc.save()
        
        logger.info(f"[OK] Completed: {doc.title} ({chunk_count} chunks)")
        return {"status": "success", "chunks": chunk_count}
        
    except Document.DoesNotExist:
        logger.error(f"Document not found: {document_id}")
        return {"status": "error", "message": "Document not found"}
    
    except ValueError as e:
        # Don't retry for content extraction errors
        error_msg = str(e)
        logger.warning(f"Document processing skipped: {error_msg}")
        
        try:
            doc = Document.objects.get(id=document_id)
            doc.status = Document.Status.FAILED
            doc.error_message = error_msg[:500]
            doc.save()
        except Exception:
            pass
        
        return {"status": "failed", "error": error_msg}
        
    except Exception as e:
        error_msg = str(e)
        logger.error(f"Document processing error: {error_msg}", exc_info=True)
        
        # Update document status
        try:
            doc = Document.objects.get(id=document_id)
            doc.status = Document.Status.FAILED
            doc.error_message = error_msg[:500]
            doc.save()
        except Exception:
            pass
        
        # Retry with exponential backoff
        retry_count = self.request.retries
        countdown = 60 * (2 ** retry_count)  # 60s, 120s, 240s
        
        logger.info(f"Retrying in {countdown}s (attempt {retry_count + 1}/3)")
        raise self.retry(exc=e, countdown=countdown)